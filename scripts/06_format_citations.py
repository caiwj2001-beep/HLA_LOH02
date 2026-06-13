#!/usr/bin/env python3
"""
HLA_LOH02 — Citation Superscript Formatter
Converts bracket citations [1], [2,3], [6-8] to Nature Medicine superscript format
"""

import re
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# ── Unicode superscript mapping ──
SUPERSCRIPT_MAP = {
    '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
    '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
    ',': '˒', '-': '⁻',  # en-dash style for ranges
}

def num_to_superscript(num_str):
    """Convert a number string to superscript Unicode."""
    return ''.join(SUPERSCRIPT_MAP.get(c, c) for c in num_str)

def compress_ranges(nums):
    """Nature style: compress consecutive numbers with hyphen.
    [1,2,3,5,6] → '1-3,5,6' """
    if len(nums) <= 2:
        return ','.join(str(n) for n in nums)
    
    nums_sorted = sorted(nums)
    result = []
    i = 0
    while i < len(nums_sorted):
        start = nums_sorted[i]
        while i + 1 < len(nums_sorted) and nums_sorted[i + 1] == nums_sorted[i] + 1:
            i += 1
        end = nums_sorted[i]
        if start == end:
            result.append(str(start))
        elif end == start + 1:
            result.append(str(start))
            result.append(str(end))
        else:
            result.append(f"{start}-{end}")
        i += 1
    
    return ','.join(result)

def convert_citation(match):
    """Convert a bracketed citation to superscript."""
    content = match.group(1).strip()
    
    # Parse all numbers
    numbers = []
    parts = content.split(',')
    for part in parts:
        part = part.strip()
        if '-' in part:
            try:
                a, b = part.split('-', 1)
                numbers.extend(range(int(a.strip()), int(b.strip()) + 1))
            except:
                numbers.append(int(part.replace('-', '').strip()))
        else:
            try:
                numbers.append(int(part))
            except:
                return match.group(0)  # Return unchanged if parsing fails
    
    if not numbers:
        return match.group(0)
    
    # Compress to Nature style
    compressed = compress_ranges(numbers)
    
    # Convert to superscript
    superscript_text = ''.join(
        num_to_superscript(c) if c.isdigit() else c 
        for c in compressed
    )
    
    return superscript_text

def process_document(input_path, output_path):
    """Process the docx file and convert all citations to superscript."""
    doc = Document(input_path)
    
    # Citation pattern: [number, comma, hyphen sequences]
    # Must be bracketed, content is only digits, commas, spaces, hyphens
    citation_pattern = re.compile(r'\[([\d,\s\-]+)\]')
    
    modified_count = 0
    
    for paragraph in doc.paragraphs:
        # Combine all runs' text to find citation positions
        full_text = paragraph.text
        matches = list(citation_pattern.finditer(full_text))
        
        if not matches:
            continue
        
        # We need to work at the run level
        runs = paragraph.runs
        
        # Strategy: process each run's text, finding and replacing citations
        for run in runs:
            text = run.text
            if not text or '[' not in text:
                continue
            
            new_text = text
            run_matches = list(citation_pattern.finditer(text))
            
            if not run_matches:
                continue
            
            # Replace from right to left to preserve positions
            for m in reversed(run_matches):
                sup_text = convert_citation(m)
                new_text = new_text[:m.start()] + sup_text + new_text[m.end():]
                modified_count += 1
            
            if new_text != text:
                run.text = new_text
    
    # Also process headers, footers, etc.
    for section in doc.sections:
        for para in section.header.paragraphs + section.footer.paragraphs:
            for run in para.runs:
                text = run.text
                if not text or '[' not in text:
                    continue
                new_text = text
                for m in reversed(list(citation_pattern.finditer(text))):
                    sup_text = convert_citation(m)
                    new_text = new_text[:m.start()] + sup_text + new_text[m.end():]
                    modified_count += 1
                if new_text != text:
                    run.text = new_text
    
    doc.save(output_path)
    print(f"Modified {modified_count} citation instances")
    print(f"Saved: {output_path}")
    return modified_count

if __name__ == '__main__':
    input_file = "/home/caiwj2001/HLA_LOH02/manuscript/HLA_LOH02_FINAL_v2.docx"
    output_file = "/home/caiwj2001/HLA_LOH02/manuscript/HLA_LOH02_FINAL_v2.docx"
    
    count = process_document(input_file, output_file)
    
    # Verify a few examples
    doc = Document(output_file)
    examples_found = 0
    for p in doc.paragraphs:
        text = p.text
        if '¹' in text and examples_found < 5:
            # Find the context around the superscript
            idx = text.index('¹')
            start = max(0, idx - 30)
            end = min(len(text), idx + 30)
            print(f"  Example: ...{text[start:end]}...")
            examples_found += 1
    
    # Check for remaining brackets
    remaining = 0
    for p in doc.paragraphs:
        if re.search(r'\[\d', p.text):
            remaining += 1
    print(f"Paragraphs with remaining [number patterns: {remaining}")
